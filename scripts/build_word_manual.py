# -*- coding: utf-8 -*-
"""WAC 창고 업무 메뉴얼 — 인턴용 하루 흐름 중심."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"C:\Users\82103\Desktop\wac 메뉴얼 만들기")
LABELS = ROOT / "wac-manual-app" / "public" / "labels"
OUT = ROOT / "WAC_창고_업무_메뉴얼.docx"
OUT_ALT = ROOT / "WAC_창고_업무_메뉴얼_v2.docx"

ACCENT = RGBColor(180, 70, 20)
MUTED = RGBColor(110, 110, 110)
DARK = RGBColor(28, 28, 28)


def set_run_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def P(doc, text, size=11, bold=False, after=6, before=0, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing = 1.3
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color or DARK)
    return p


def title_block(doc, no, title, sub):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(no), size=11, bold=True, color=ACCENT)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    set_run_font(p2.add_run(title), size=17, bold=True, color=DARK)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(10)
    set_run_font(p3.add_run(sub), size=10, color=MUTED)


def bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_run_font(p.add_run(text), size=size)


def numbered(doc, text, size=11):
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    set_run_font(p.add_run(text), size=size)


def img(doc, path: Path, width_cm: float, caption=None):
    if not path.exists():
        P(doc, f"[이미지 없음: {path.name}]", size=10, color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(6)
        set_run_font(c.add_run(caption), size=9, color=MUTED)


def page_break(doc):
    doc.add_page_break()


def two_col_img_text(doc, image_path: Path, img_w: float, lines: list[str], caption: str):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left, right = table.rows[0].cells
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path.exists():
        lp.add_run().add_picture(str(image_path), width=Cm(img_w))
    cap = left.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    set_run_font(cap.add_run(caption), size=8, color=MUTED)
    right.paragraphs[0].clear()
    for i, text in enumerate(lines):
        p = right.paragraphs[0] if i == 0 else right.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.line_spacing = 1.2
        set_run_font(p.add_run(text), size=10)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.9)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    # ===== COVER =====
    P(doc, "", after=40)
    P(doc, "WAC 창고 업무 메뉴얼", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    P(doc, "인턴 · 신규 직원용 피킹 / 패킹 가이드", size=13, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    P(doc, "W MART  ·  A Member of WAC", size=11, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    P(
        doc,
        "패킹리스트 분류부터 피킹·패킹·라벨까지 순서대로 따라가면 됩니다.",
        size=11,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
    )

    # ===== TOC =====
    page_break(doc)
    P(doc, "목차", size=18, bold=True, after=16)
    P(doc, "1.  창고 작업 수칙", size=13, bold=True, after=4)
    P(doc, "     분류 · 피킹 · 손수레 · 박스 패킹 · 라벨", size=10, color=MUTED, after=12)
    P(doc, "2.  창고 도면", size=13, bold=True, after=4)
    P(doc, "     전체 도면 · 렉 단 사진", size=10, color=MUTED, after=8)

    # ===== RULE 01 =====
    page_break(doc)
    title_block(doc, "수칙 01", "패킹리스트 분류 · 피킹 순서", "출력 후 → 보고 나누기 → 피킹")
    P(doc, "패킹리스트 출력 후, 인턴이 할 일", size=11, bold=True, after=4)
    P(
        doc,
        "거래처코드 맨 앞 글자(K / N / H)와 거래처명을 보고 아래 순서로 분류합니다.",
        size=10,
        after=8,
    )
    P(doc, "K · N · H = 배송 지역", size=11, bold=True, after=4)
    bullet(doc, "K — Kowloon (구룡)")
    bullet(doc, "N — New Territories (신계)")
    bullet(doc, "H — Hong Kong Island (홍콩섬)")
    P(doc, "", after=4)
    P(doc, "보는 법 1) 거래처코드 맨 앞 = K / N / H", size=10, bold=True, after=2)
    P(doc, "보는 법 2) K이면 거래처명에 ‘침사추이’ 있는지 확인 → 있으면 맨 먼저 분류", size=10, bold=True, after=10)

    P(doc, "예시 ①  K + 침사추이  (가장 먼저)", size=11, bold=True, after=4)
    img(doc, LABELS / "example-K-chimsa.png", width_cm=15.2, caption="①K + ①침사추이 · 판매조회 실제 패킹리스트")

    P(doc, "예시 ②  N", size=11, bold=True, after=4)
    img(doc, LABELS / "example-N.png", width_cm=13.5, caption="N · 거래처코드 맨 앞 글자만")

    P(doc, "분류 · 피킹 순서", size=11, bold=True, after=4)
    numbered(doc, "K · 침사추이  (맨 먼저 분류 · 예시 ①)")
    numbered(doc, "K 나머지  (K인데 침사추이 아님)")
    numbered(doc, "N  (예시 ②)")
    numbered(doc, "H  → 끝나면 패킹")
    P(doc, "K 침사추이  →  K  →  N  →  H  →  패킹", size=12, bold=True, after=4)

    # ===== RULE 02 =====
    page_break(doc)
    title_block(doc, "수칙 02", "패킹리스트로 피킹 · 패킹하기", "이니셜 · 위치 · 수량 · 박스 라벨")
    steps = [
        "① Picked by — 영어 이니셜\n(예: 옥현서 → OK)",
        "② 위치 — 렉에서 찾기\n(예: A 5/6)",
        "③ 수량만큼 피킹\n(예: 10PK)",
        "패킹하는 곳 — 바닥에 내려놓고 패킹",
        "④ 박스수량 기입 (예: 3)",
        "⑤ 거래처코드 라벨\n(예: 222-02×3)",
    ]
    two_col_img_text(
        doc,
        LABELS / "packing-list-annotated.png",
        8.2,
        steps,
        "빨간 원 = 작업 포인트",
    )

    # ===== RULE 03 — cart =====
    page_break(doc)
    title_block(doc, "수칙 03", "파란 손수레 쌓기 · 내리기", "앞부터 쌓기 / 뒤부터 내리기")
    P(doc, "손잡이 기준으로 앞과 뒤가 반대입니다.", size=11, after=8)
    bullet(doc, "피킹(쌓기) — 손잡이에서 먼 앞쪽부터")
    bullet(doc, "패킹 전(내리기) — 손잡이 쪽 뒤쪽부터")
    img(doc, LABELS / "cart-annotated.png", width_cm=8.8, caption="① 뒤(손잡이)  /  ② 앞")

    # ===== RULE 04 — box packing at packing area =====
    page_break(doc)
    title_block(doc, "수칙 04", "박스에 넣을 때 (패킹)", "더블체크 → 온도대 · 칸막이")
    P(
        doc,
        "피킹이 끝나면 패킹하는 곳에서 박스에 담습니다.\n"
        "단, Checked by에 영어 이니셜이 있는 리스트만 패킹합니다. (더블체크 완료)",
        size=10,
        after=8,
    )
    P(doc, "더블체크 확인 (패킹 전에 꼭)", size=11, bold=True, after=4)
    P(
        doc,
        "Picked by 옆 Checked by에 이니셜이 있으면 → 더블체크 완료 → 패킹 가능",
        size=10,
        after=6,
    )
    img(
        doc,
        LABELS / "packing-list-doublecheck.png",
        width_cm=12.5,
        caption="Checked by에 이니셜이 있으면 패킹",
    )
    P(doc, "박스에 넣는 규칙", size=11, bold=True, after=6)
    P(doc, "1.  온도대별 분리 포장", size=11, bold=True, after=2)
    P(doc, "상온과 냉장·냉동은 원칙적으로 다른 박스에 넣습니다.", size=10, after=8)
    P(doc, "2.  합포장(혼재)할 때", size=11, bold=True, after=2)
    P(
        doc,
        "어쩔 수 없이 한 박스에 섞으면, 종이 칸막이(파티션)로\n제품이 서로 닿지 않게 합니다.",
        size=10,
        after=8,
    )
    P(doc, "3.  분말류 방습", size=11, bold=True, after=2)
    P(doc, "분말형 제품은 가급적 단독으로 밀봉 포장합니다.", size=10, after=8)
    P(doc, "4.  깻잎 · 청양고추 — 혼합 금지", size=11, bold=True, after=2)
    P(doc, "깻잎, 청양고추는 냉장·냉동 제품과 같은 박스에 넣지 않습니다.", size=10, after=10)
    P(
        doc,
        "기억: 더블체크 → 온도 분리 → (섞으면) 칸막이 → 분말은 단독 → 깻잎·청양고추는 냉장·냉동과 금지",
        size=10,
        bold=True,
        after=4,
    )

    # ===== RULE 05 — box type by weight =====
    page_break(doc)
    title_block(doc, "수칙 05", "박스 고르기 — 무게별", "무거우면 두면 · 가벼우면 한면")
    P(
        doc,
        "패킹할 때 박스가 무거워질 것 같으면 단면이 물결 두 줄(두면)인 박스를 씁니다.\n"
        "가벼운 짐이면 물결 한 줄(한면)이면 됩니다. 모서리 단면을 보고 고르세요.",
        size=10,
        after=8,
    )
    P(doc, "무거운 짐 → 두면 (두 줄)", size=11, bold=True, after=4)
    P(doc, "대략 10kg 초과로 무거울 때. 박스 예시:", size=10, after=2)
    P(doc, "· 삼계탕 박스", size=10, after=1)
    P(doc, "· 물결가루(치킨파우더) 박스", size=10, after=1)
    P(doc, "· W 자체 박스", size=10, after=4)
    img(doc, LABELS / "box-double-wall.png", width_cm=12.0, caption="빨간 원 = 단면 물결 두 줄")
    P(doc, "가벼운 짐 → 한면 (한 줄)", size=11, bold=True, after=4)
    P(doc, "가벼운 짐일 때. 박스 예시:", size=10, after=2)
    P(doc, "· 김가루 박스 (예: 해농 넘버원)", size=10, after=1)
    P(doc, "· 떡볶이(떡) 박스", size=10, after=4)
    img(doc, LABELS / "box-single-wall.png", width_cm=12.0, caption="빨간 원 = 단면 물결 한 줄")
    P(doc, "한 줄: 무거우면 두면 · 가벼우면 한면", size=11, bold=True, after=4)

    # ===== RULE 06 — caution + heavy SAME PAGE =====
    page_break(doc)
    title_block(doc, "수칙 06", "라벨 붙이기 — 주의 · HEAVY", "터짐 주의 / 10kg 초과")
    duo = doc.add_table(rows=1, cols=2)
    duo.autofit = True
    c1, c2 = duo.rows[0].cells

    c1.paragraphs[0].clear()
    p = c1.paragraphs[0]
    set_run_font(p.add_run("주의 (FRAGILE)"), size=12, bold=True)
    p2 = c1.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    set_run_font(
        p2.add_run("단무지, 마요네즈, 깐메추리알, 케찹팩 등\n비닐로 감싼 뒤 주의 라벨"),
        size=9,
    )
    if (LABELS / "caution.png").exists():
        ip = c1.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(str(LABELS / "caution.png"), width=Cm(6.8))

    c2.paragraphs[0].clear()
    p = c2.paragraphs[0]
    set_run_font(p.add_run("HEAVY"), size=12, bold=True)
    p2 = c2.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    set_run_font(
        p2.add_run("패킹 후 박스가 무거우면 HEAVY 라벨\n(대략 10kg 초과)"),
        size=9,
    )
    if (LABELS / "heavy.png").exists():
        ip = c2.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(str(LABELS / "heavy.png"), width=Cm(6.8))

    doc.add_paragraph()

    # ===== MAP 01 =====
    page_break(doc)
    title_block(doc, "도면 01", "창고 전체 도면", "왼쪽 A · 가운데 통로 · 오른쪽 B")
    img(doc, LABELS / "warehouse-map.png", width_cm=14.5)

    # ===== MAP 02 =====
    page_break(doc)
    title_block(doc, "도면 02", "렉 단(높이) 사진", "위 → 아래")
    P(doc, "1/3 (상단)  →  1/2 (중단)  →  1/1 (하단·팔레트)", size=11, bold=True, after=8)
    annotated = LABELS / "rack-levels-annotated.png"
    photo = LABELS / "rack-levels.png"
    img(
        doc,
        annotated if annotated.exists() else photo,
        width_cm=11.0,
        caption="실제 렉 · 상단 1/3 · 중단 1/2 · 하단 1/1",
    )
    bullet(doc, "1/3 — 작은 박스 · 가벼운 짐")
    bullet(doc, "1/2 — 중간 크기 · 오픈픽")
    bullet(doc, "1/1 — 바닥 팔레트 · 큰/무거운 짐")

    try:
        doc.save(OUT)
        print("saved", OUT)
    except PermissionError:
        doc.save(OUT_ALT)
        print("원본 열림 →", OUT_ALT)


if __name__ == "__main__":
    build()
