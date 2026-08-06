# -*- coding: utf-8 -*-
"""다음 인턴용 — 메뉴얼 수정 가이드 Word 생성."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"C:\Users\82103\Desktop\wac 메뉴얼 만들기")
OUT = ROOT / "메뉴얼_수정_가이드_인턴용.docx"

DARK = RGBColor(28, 28, 28)
MUTED = RGBColor(100, 100, 100)
ACCENT = RGBColor(180, 70, 20)


def font(run, size=11, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def P(doc, text, size=11, bold=False, after=6, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run(text)
    font(r, size=size, bold=bold, color=color or DARK)
    return p


def H(doc, text, size=16):
    P(doc, text, size=size, bold=True, after=10)


def bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    font(r, size=size)


def numbered(doc, text, size=11):
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    font(r, size=size)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    P(doc, "WAC 창고 메뉴얼 — 수정 가이드", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    P(doc, "다음 인턴·담당자용 (쉽게 고치는 방법)", size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

    H(doc, "1. 이게 뭔가요?")
    P(
        doc,
        "창고 업무 전자 메뉴얼 웹사이트입니다. 주소는 아래와 같습니다.",
        after=4,
    )
    P(doc, "https://wac-warehouse-manual.vercel.app", size=11, bold=True, after=8)
    P(doc, "코드·사진은 GitHub에 공개되어 있습니다.", after=4)
    P(doc, "https://github.com/hyunseook0606-dev/wac-manual", size=11, bold=True, after=12)

    H(doc, "2. 어디에 무엇이 있나요?")
    bullet(doc, "src/components/RulesChapter.tsx  →  작업 수칙 글·순서 (가장 자주 고침)")
    bullet(doc, "src/components/MapChapter.tsx  →  창고 도면 화면")
    bullet(doc, "src/components/CoverChapter.tsx  →  시작 화면")
    bullet(doc, "public/labels/  →  사진·라벨 이미지 (png)")
    bullet(doc, "scripts/build_word_manual.py  →  인쇄용 Word 메뉴얼 다시 만들기")
    P(doc, "", after=4)

    H(doc, "3. 글만 고치고 싶을 때 (제일 흔함)")
    numbered(doc, "GitHub에서 wac-manual 저장소를 연다")
    numbered(doc, "src/components/RulesChapter.tsx 파일을 연다")
    numbered(doc, "원하는 문장을 고친다 (규칙 제목·설명·순서)")
    numbered(doc, "Commit changes → main에 저장")
    numbered(doc, "1~2분 후 웹사이트(Vercel)에 자동 반영되는지 확인")
    P(
        doc,
        "Tip: RULE_PAGES 배열 안에 규칙이 페이지 단위로 들어 있습니다. 새 규칙은 그 배열에 항목을 추가하면 됩니다.",
        size=10,
        color=MUTED,
        after=12,
    )

    H(doc, "4. 사진을 바꾸고 싶을 때")
    numbered(doc, "새 사진을 public/labels/ 폴더에 넣는다 (예: my-photo.png)")
    numbered(doc, "RulesChapter.tsx에서 이미지 경로를 /labels/my-photo.png 로 바꾼다")
    numbered(doc, "GitHub에 커밋하면 웹에 반영된다")
    P(doc, "", after=4)

    H(doc, "5. 인쇄용 Word도 같이 고치려면")
    numbered(doc, "scripts/build_word_manual.py 내용도 웹과 맞게 수정")
    numbered(doc, "컴퓨터에 Python + python-docx 설치 후 실행:")
    P(doc, "python scripts/build_word_manual.py", size=10, bold=True, after=6)
    numbered(doc, "만들어진 WAC_창고_업무_메뉴얼.docx 를 확인·배포")
    P(doc, "", after=4)

    H(doc, "6. 로컬에서 미리보기 (선택)")
    numbered(doc, "저장소를 내 컴퓨터에 Clone")
    numbered(doc, "폴더에서 npm install")
    numbered(doc, "npm run dev 실행 후 브라우저에서 확인")
    P(doc, "", after=4)

    H(doc, "7. 막히면")
    bullet(doc, "GitHub 저장소 Issues에 메모를 남기거나")
    bullet(doc, "Cursor / VS Code로 폴더를 열고 RulesChapter.tsx만 검색해서 수정")
    bullet(doc, "Vercel 대시보드에서 최근 Deploy가 성공했는지 확인")
    P(doc, "", after=8)

    P(
        doc,
        "핵심 한 줄: 글·사진은 GitHub에서 고치면 → Vercel 웹이 따라 업데이트됩니다.",
        size=12,
        bold=True,
        after=4,
    )

    doc.save(OUT)
    print("saved", OUT)


if __name__ == "__main__":
    build()
